#!/usr/bin/env python3
"""Multilingual search test — verifies peekdocs finds text in 14 languages.

Run from the repo root with the venv activated:
    source venv/bin/activate
    python tests/language_test.py

Creates temporary test files, runs peekdocs against each language,
and prints a summary table. Cleans up after itself.
"""

import os
import subprocess
import shutil
import sys

from docx import Document


def main():
    test_dir = "/tmp/peekdocs_language_test"
    if os.path.exists(test_dir):
        shutil.rmtree(test_dir)
    os.makedirs(test_dir)

    # Create multilingual .txt file
    with open(os.path.join(test_dir, "multilingual.txt"), "w", encoding="utf-8") as f:
        f.write(
            "English: The budget report is ready for review.\n"
            "Chinese: 预算报告已准备好供审查。\n"
            "Japanese: 予算報告書はレビューの準備ができました。\n"
            "Korean: 예산 보고서가 검토 준비가 되었습니다.\n"
            "Arabic: تقرير الميزانية جاهز للمراجعة.\n"
            "Hindi: बजट रिपोर्ट समीक्षा के लिए तैयार है।\n"
            "Russian: Бюджетный отчёт готов к проверке.\n"
            "Spanish: El informe presupuestario está listo para su revisión.\n"
            "German: Der Haushaltsbericht ist zur Überprüfung bereit.\n"
            "French: Le rapport budgétaire est prêt pour examen.\n"
            "Portuguese: O relatório orçamentário está pronto para revisão.\n"
            "Thai: รายงานงบประมาณพร้อมสำหรับการตรวจสอบ\n"
            "Hebrew: דוח התקציב מוכן לבדיקה.\n"
            "Greek: Η αναφορά προϋπολογισμού είναι έτοιμη για έλεγχο.\n"
        )

    # Create Chinese .docx
    doc = Document()
    doc.add_paragraph("公司财务报告 2025年第一季度")
    doc.add_paragraph("总预算：¥500,000")
    doc.save(os.path.join(test_dir, "chinese_report.docx"))

    # Create Greek .docx
    doc = Document()
    doc.add_paragraph("Τριμηνιαία Οικονομική Αναφορά 2025")
    doc.add_paragraph("Συνολικός προϋπολογισμός: €500.000")
    doc.save(os.path.join(test_dir, "greek_report.docx"))

    languages = [
        ("English",    "budget",           "Latin"),
        ("Chinese",    "预算",              "Hanzi"),
        ("Japanese",   "予算",              "Kanji/Kana"),
        ("Korean",     "예산",              "Hangul"),
        ("Arabic",     "الميزانية",         "Arabic"),
        ("Hindi",      "बजट",              "Devanagari"),
        ("Russian",    "Бюджетный",        "Cyrillic"),
        ("Greek",      "προϋπολογισμού",   "Greek"),
        ("Spanish",    "presupuestario",   "Latin"),
        ("German",     "Haushaltsbericht", "Latin"),
        ("French",     "budgétaire",       "Latin"),
        ("Portuguese", "orçamentário",     "Latin"),
        ("Thai",       "งบประมาณ",          "Thai"),
        ("Hebrew",     "התקציב",            "Hebrew"),
    ]

    passed = 0
    failed = 0
    results = []

    for lang, term, script in languages:
        r = subprocess.run(
            ["peekdocs", term],
            capture_output=True, text=True, cwd=test_dir, timeout=30,
        )
        ok = r.returncode == 0
        if ok:
            passed += 1
        else:
            failed += 1
        results.append((lang, script, term, ok))

    # Print table
    print()
    print("  Language       Script          Search Term              Result")
    print("  ───────────── ─────────────── ──────────────────────── ──────")
    for lang, script, term, ok in results:
        status = "PASS" if ok else "FAIL"
        print(f"  {lang:13s}  {script:15s}  {term:24s}  {status}")

    print()
    print(f"  {passed + failed} languages tested  |  {passed} passed  |  {failed} failed")
    print(f"  File types: .txt, .docx")
    print(f"  Scripts: Latin, CJK, Arabic, Devanagari, Cyrillic, Greek, Thai, Hebrew")
    print()

    if failed == 0:
        print("  ALL LANGUAGES VERIFIED")
    else:
        print(f"  {failed} LANGUAGE(S) FAILED")

    # Clean up
    shutil.rmtree(test_dir)

    return 0 if failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
