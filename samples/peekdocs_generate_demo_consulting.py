"""Generate the peekdocs_demo_consulting/ folder from scratch.

Run from anywhere with peekdocs installed (its python-docx and
fpdf2 dependencies do the heavy lifting). Reproducible — re-run
to regenerate; binary outputs are intentionally not committed
to git.

Output: 6 .docx files + 2 .pdf files in
samples/peekdocs_demo_consulting/, each hand-tuned to land hits
for the Quarterly Content Audit suite (Project Apex / Acme Corp /
Legacy Reports / status flags / dollar amounts / Q2 2026
references).
"""

import os
from docx import Document
from fpdf import FPDF

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "peekdocs_demo_consulting")
os.makedirs(OUT, exist_ok=True)


def write_docx(name, paragraphs, heading=None):
    """Write a .docx with an optional heading and a list of paragraphs."""
    doc = Document()
    if heading:
        doc.add_heading(heading, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(os.path.join(OUT, name))


def write_pdf(name, lines, title=None):
    """Write a single-page .pdf with optional title and a list of body lines."""
    pdf = FPDF()
    pdf.add_page()
    if title:
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, title, ln=True)
        pdf.ln(4)
    pdf.set_font("Helvetica", size=11)
    for line in lines:
        pdf.multi_cell(0, 6, line)
        pdf.ln(1)
    pdf.output(os.path.join(OUT, name))


# 1. Project Apex Q2 status report
write_docx(
    "Q2_2026_status_project_apex.docx",
    heading="Project Apex — Q2 2026 Status Report",
    paragraphs=[
        "Reporting period: 2026-04-01 through 2026-06-30 (Q2 2026).",
        "Project Apex is currently at risk. Two of the three milestone targets "
        "for Q2 2026 have slipped: the API freeze (originally 2026-05-15) is "
        "blocked pending the legal review of the Acme Corp data-sharing "
        "agreement, and the migration off Legacy Reports has been delayed by "
        "approximately three weeks.",
        "Forecast spend for Q2 2026: $245,000. Variance vs. plan: +$32,000 "
        "(see budget summary).",
        "Mitigation: weekly Project Apex sync with Acme Corp account team "
        "starting 2026-06-20, plus a dedicated push to retire the last "
        "Legacy Reports consumer by 2026-06-30.",
    ],
)

# 2. Acme Corp statement of work
write_docx(
    "acme_corp_statement_of_work.docx",
    heading="Acme Corp — Statement of Work",
    paragraphs=[
        "Engagement: Project Apex integration support.",
        "Period: 2026-04-15 through 2026-12-31.",
        "Total committed fees: $180,000, billed monthly at $20,000 per month "
        "through Q2 2026 and Q3 2026, then $15,000 per month through Q4 2026.",
        "Status as of 2026-06-15: at risk. Acme Corp legal has flagged two "
        "clauses related to the Legacy Reports data export pipeline; resolution "
        "blocked pending counsel review.",
        "Next review: 2026-06-22.",
    ],
)

# 3. Legacy Reports migration plan
write_docx(
    "migration_plan_legacy_reports.docx",
    heading="Migration Plan: Legacy Reports → New Reporting Platform",
    paragraphs=[
        "Owner: Project Apex working group.",
        "Target sunset date for Legacy Reports: 2026-06-30 (end of Q2 2026). "
        "Status: blocked — the Acme Corp consumer hasn't yet validated the new "
        "platform's export format.",
        "Risk: if Legacy Reports remains live past 2026-06-30, the parallel-run "
        "cost ($8,500/month) continues into Q3 2026.",
        "Open items: (a) Acme Corp signoff on format, (b) decommission the last "
        "two reports still consumed externally, (c) archive read-only snapshots.",
    ],
)

# 4. Q2 2026 budget summary
write_docx(
    "q2_2026_budget_summary.docx",
    heading="Q2 2026 Budget Summary",
    paragraphs=[
        "Period: 2026-04-01 through 2026-06-30.",
        "Total committed Q2 2026 spend: $612,000.",
        "Largest line items: Project Apex engineering ($245,000), Acme Corp "
        "fees ($60,000 across April / May / June), Legacy Reports parallel-run "
        "($25,500), cloud infrastructure ($118,000), contractors ($75,000).",
        "Variance: +$48,000 over plan. Driver: Project Apex slippage forcing "
        "an extra month of Legacy Reports parallel-run cost.",
        "Q3 2026 forecast pending the 2026-07-05 planning session.",
    ],
)

# 5. Weekly status — heavy on status flags
write_docx(
    "weekly_status_2026_05_15.docx",
    heading="Weekly Status — Week of 2026-05-15",
    paragraphs=[
        "Project Apex: at risk. Two milestones slipped, blocked on Acme Corp "
        "legal review.",
        "Legacy Reports migration: blocked. Awaiting Acme Corp signoff on "
        "export format.",
        "Q2 2026 budget: tracking +$32,000 over plan; under review.",
        "Vendor X integration: delayed to 2026-06-08; minor impact.",
        "Hiring (Apex backend lead): at risk. Top candidate withdrew; restarting "
        "search.",
        "All other workstreams: green.",
    ],
)

# 6. Acme Corp risks register
write_docx(
    "acme_corp_risks_register.docx",
    heading="Acme Corp Engagement — Risks Register",
    paragraphs=[
        "Last updated 2026-06-10.",
        "R1: Acme Corp legal review of the Project Apex SOW blocked. Owner: "
        "Legal. Exposure: $45,000 (one month of Acme Corp fees at risk if the "
        "engagement pauses).",
        "R2: Legacy Reports format incompatibility with the new platform. "
        "Owner: Platform team. Status: blocked, mitigation in progress.",
        "R3: Q2 2026 spend tracking over plan by $48,000. Owner: Finance. "
        "Status: at risk, may force Q3 2026 scope cut.",
        "R4: Acme Corp turnover on their integration team — two of three "
        "engineers departed in May 2026. Status: delayed onboarding of "
        "replacements until 2026-06-25.",
    ],
)

# 7. Quarterly roadmap (PDF)
write_pdf(
    "2026_q2_roadmap.pdf",
    title="2026 Q2 Roadmap",
    lines=[
        "Period: 2026-04-01 through 2026-06-30.",
        "",
        "Project Apex (priority 1):",
        "  - Milestone A: API freeze. Target 2026-05-15. Status: blocked.",
        "  - Milestone B: Acme Corp signoff. Target 2026-06-01. Status: at risk.",
        "  - Milestone C: Beta launch. Target 2026-06-30. Status: at risk.",
        "",
        "Legacy Reports retirement (priority 2):",
        "  - Sunset target: 2026-06-30. Status: blocked.",
        "  - Parallel-run cost if delayed: $8,500/month.",
        "",
        "Q2 2026 spend forecast: $612,000 (variance +$48,000).",
        "Next quarterly review: 2026-07-05.",
    ],
)

# 8. Acme Corp invoice (PDF)
write_pdf(
    "acme_corp_invoice_2026_05.pdf",
    title="Invoice INV-2026-05-0118",
    lines=[
        "Bill to: Acme Corp",
        "Period: 2026-05-01 through 2026-05-31",
        "",
        "Project Apex consulting (May 2026):           $20,000.00",
        "Legacy Reports parallel-run support (May):     $4,250.00",
        "Travel and expenses (Q2 2026 portion):           $850.00",
        "",
        "Subtotal:                                     $25,100.00",
        "Tax:                                           $1,757.00",
        "Total due:                                    $26,857.00",
        "",
        "Due date: 2026-06-30. Net 30 from invoice date.",
    ],
)

print(f"Generated 6 .docx + 2 .pdf in {OUT}")
print("Files:")
for f in sorted(os.listdir(OUT)):
    print(f"  {f}")
