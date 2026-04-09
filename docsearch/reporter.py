"""Report generation for docsearch (TXT, DOCX, CSV, JSON, append, suite reports)."""

import csv
import json
import os
import re
import shutil
import textwrap
from copy import deepcopy
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor, Inches

from docsearch.scanner import _wildcard_to_regex
from docsearch.translator import translate_search


def _safe_filename_part(name):
    """Sanitize a name for use in a filename (replace non-word chars with _)."""
    return re.sub(r'[^\w\-]', '_', name)


def copy_stage_reports(results_dir, suite_name, stage_index, search_name, timestamp_suffix=""):
    """Copy docsearch_results.* to named stage files (DO_NOT_SEARCH_ prefixed).

    Returns a dict mapping extension to the stage file path.
    """
    safe_suite = _safe_filename_part(suite_name)
    safe_search = _safe_filename_part(search_name)
    ts = f"_{timestamp_suffix}" if timestamp_suffix else ""
    prefix = f"DO_NOT_SEARCH_SUITE_{safe_suite}_stage{stage_index:02d}_{safe_search}{ts}"

    copied = {}
    for ext in ("txt", "docx", "csv", "json"):
        src = os.path.join(results_dir, f"docsearch_results.{ext}")
        if os.path.exists(src):
            dst = os.path.join(results_dir, f"{prefix}.{ext}")
            shutil.copy2(src, dst)
            copied[ext] = dst
    return copied


def cleanup_stage_reports(results_dir, suite_name):
    """Remove any existing stage report files for the given suite name."""
    safe_suite = _safe_filename_part(suite_name)
    pattern = f"DO_NOT_SEARCH_SUITE_{safe_suite}_stage"
    for fname in os.listdir(results_dir):
        if fname.startswith(pattern):
            os.remove(os.path.join(results_dir, fname))


def fmt_size(b):
    """Format byte count as human-readable string."""
    if b >= 1_000_000:
        return f"{b / 1_000_000:.2f} MB"
    elif b >= 1_000:
        return f"{b / 1_000:.2f} KB"
    return f"{b} bytes"


def _strip_highlights(text):
    """Remove ** highlight markers from text."""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def write_txt_report(output_path, matches, all_files, search_terms, command_str,
                     report_mode, use_ocr, exclude_terms, use_context,
                     use_fuzzy, use_regex, use_wildcard,
                     search_elapsed, cores, cpu_count,
                     inverse_files=None, recursive=False,
                     file_types=None, proximity=None,
                     context_before=0, context_after=0,
                     specific_files=None, use_index=False,
                     inverse=False, output_csv=False, output_json=False,
                     expression=None, use_whole_word=False,
                     total_matches=None, max_matches=None,
                     range_specs=None, index_meta=None):
    """Write docsearch_results.txt report file.

    Returns (total_bytes, size_str) for use in console summary.
    """
    if os.path.exists(output_path):
        os.remove(output_path)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("Header:\n\n")
        f.write("Program name: docsearch\n")
        f.write("Program Source: https://github.com/exbuf\n")
        f.write("Overview: Searches all supported file types in current directory for search terms.\n")
        f.write("Supported file types:\n")
        f.write(".7z, .bz2, .cfg, .csv, .doc, .docx, .eml, .epub, .gz, .html, .ini, .json, .log, .md, .msg,\n")
        f.write(".ods, .odp, .odt, .pdf, .ppt, .pptx, .pst, .rar, .rst, .rtf, .sql, .tar, .tex, .tgz, .toml,\n")
        f.write(".tsv, .txt, .xls, .xlsx, .xml, .yaml, .yml, .zip\n")
        if use_ocr:
            f.write("OCR image types: .bmp, .jpg, .jpeg, .png, .tif, .tiff\n")
        f.write(f"\nReport Generated On ==> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Command ==> {command_str}\n")
        translation = translate_search(
            search_terms, report_mode=report_mode,
            use_regex=use_regex, use_wildcard=use_wildcard,
            use_fuzzy=use_fuzzy, use_ocr=use_ocr, use_index=use_index,
            use_whole_word=use_whole_word,
            inverse=inverse, recursive=recursive,
            exclude_terms=exclude_terms,
            file_types=file_types, specific_files=specific_files,
            proximity=proximity,
            context_before=context_before, context_after=context_after,
            expression=expression,
            range_specs=range_specs,
        )
        f.write(f"Translation ==> {translation}\n")
        if expression:
            f.write(f"Search Expression ==> {expression} (mode: {report_mode})\n")
        else:
            f.write(f"Search Term(s) ==> {' '.join(search_terms)} (match: {report_mode})\n")
        if exclude_terms:
            f.write(f"Exclude Term(s) ==> {' '.join(exclude_terms)}\n")
        if total_matches is not None and total_matches > len(matches):
            f.write(f"Hits ==> {len(matches)} (of {total_matches:,} total — report capped at {max_matches:,}; use -m to change)\n")
        else:
            f.write(f"Hits ==> {len(matches)}\n")

        # ── Search Settings ──
        on_off = lambda v: "ON" if v else "OFF"
        f.write(f"\nSearch Settings:\n")
        f.write(f"  AND mode: {on_off(report_mode == 'ALL')}  |  Recursive: {on_off(recursive)}  |  Inverse: {on_off(inverse)}  |  Expression: {on_off(expression is not None)}\n")
        f.write(f"  Fuzzy: {on_off(use_fuzzy)}  |  Wildcard: {on_off(use_wildcard)}  |  Regex: {on_off(use_regex)}  |  Whole Word: {on_off(use_whole_word)}  |  OCR: {on_off(use_ocr)}\n")
        f.write(f"  Index: {on_off(use_index)}\n")
        if use_index and index_meta:
            f.write(f"  Index last updated: {index_meta.get('last_updated', index_meta.get('created_at', '?'))}"
                    f"  ({index_meta.get('file_count', '?')} files, {index_meta.get('line_count', '?')} lines)\n")
        if file_types:
            f.write(f"  File types: {file_types}\n")
        if specific_files:
            f.write(f"  Specific files: {specific_files}\n")
        if proximity:
            f.write(f"  Proximity: {proximity} words\n")
        if use_context:
            f.write(f"  Context: {context_before} line(s) before, {context_after} line(s) after\n")
        outputs = ["TXT", "DOCX"]
        if output_csv:
            outputs.append("CSV")
        if output_json:
            outputs.append("JSON")
        f.write(f"  Output formats: {', '.join(outputs)}\n\n")
        # Per-file match counts (used later for per-match headers)
        file_counts = {}
        for fd, fn, _ln, _tx in matches:
            key = (fd, fn)
            if key not in file_counts:
                file_counts[key] = 0
            file_counts[key] += 1
        f.write(f"Search Time ==> {search_elapsed:.2f} seconds, Cores used ==> {cores} of {cpu_count}\n")
        total_bytes = sum(os.path.getsize(f_path) for f_path in all_files)
        if total_bytes >= 1_000_000:
            size_str = f"{total_bytes / 1_000_000:.2f} MB"
        elif total_bytes >= 1_000:
            size_str = f"{total_bytes / 1_000:.2f} KB"
        else:
            size_str = f"{total_bytes} bytes"
        f.write(f"Files searched ==> {len(all_files)} ({size_str})\n")
        ext_counts = {}
        for fp in all_files:
            ext = os.path.splitext(fp)[1].lower()
            ext_counts[ext] = ext_counts.get(ext, 0) + 1
        if ext_counts:
            tally = ", ".join(f"{ext}: {count}" for ext, count in ext_counts.items())
            f.write(f"File Types Searched ==> {tally}\n")
        f.write("\n")
        f.write("Results:\n\n")
        if inverse_files is not None:
            f.write(f"Files WITHOUT matches: {len(inverse_files)} (out of {len(all_files)} searched)\n\n")
            for fp in inverse_files:
                f.write(f"  {os.path.basename(fp)}\n")
                f.write(f"  ({os.path.dirname(fp)})\n\n")
            # Inverse reports only list files missing the term — no match lines
            return (total_bytes, size_str)
        prev_file = None
        for file_dir, filename, line_num, text in matches:
            current_file = os.path.join(file_dir, filename)
            if use_context and prev_file == current_file:
                f.write("---\n\n")
            prev_file = current_file
            if use_context:
                lines = text.split("\n")
                wrapped_lines = [textwrap.fill(line, width=80) if line else line for line in lines]
                wrapped = "\n".join(wrapped_lines)
            else:
                wrapped = textwrap.fill(text, width=80)
            fc = file_counts.get((file_dir, filename), 1)
            f.write(f'Document: {filename} ({fc} match{"es" if fc != 1 else ""}), Line: {line_num}, Match:\n({file_dir})\n"{wrapped}"\n\n')

    return (total_bytes, size_str)


def write_docx_report(docx_path, txt_path, search_terms=None,
                      use_regex=False, use_wildcard=False,
                      use_whole_word=False, use_fuzzy=False,
                      expression=None):
    """Create docsearch_results.docx from the .txt report with yellow highlighting.

    Highlights matched terms directly using the search parameters rather than
    parsing markers from the text. Returns the Document object for further
    modification.
    """
    if os.path.exists(docx_path):
        os.remove(docx_path)

    # Build highlight patterns from search terms
    highlight_patterns = []
    if not use_fuzzy:
        if expression:
            from docsearch.expr_parser import parse_expression, extract_positive_terms
            highlight_terms = extract_positive_terms(parse_expression(expression))
        elif search_terms:
            highlight_terms = search_terms
        else:
            highlight_terms = []
        for term in highlight_terms:
            if use_wildcard:
                highlight_patterns.append(_wildcard_to_regex(term))
            elif use_regex:
                highlight_patterns.append(term)
            elif use_whole_word:
                highlight_patterns.append(r'\b' + re.escape(term) + r'\b')
            else:
                highlight_patterns.append(re.escape(term))

    # Combine into one pattern for efficient matching
    combined_pattern = None
    if highlight_patterns:
        try:
            combined_pattern = re.compile("|".join(highlight_patterns), re.IGNORECASE)
        except re.error:
            combined_pattern = None

    def _add_highlighted_line(para, line, bold=False):
        """Add a line to a paragraph with yellow highlighting on matched terms."""
        if combined_pattern:
            parts = combined_pattern.split(line)
            matches = combined_pattern.findall(line)
            for i, part in enumerate(parts):
                if part:
                    run = para.add_run(part)
                    if bold:
                        run.bold = True
                if i < len(matches):
                    run = para.add_run(matches[i])
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    if bold:
                        run.bold = True
        else:
            run = para.add_run(line)
            if bold:
                run.bold = True

    result_doc = Document()
    in_results = False
    with open(txt_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.rstrip("\n")
            para = result_doc.add_paragraph()

            # Make URL a clickable hyperlink
            if line.startswith("Program Source: "):
                prefix = "Program Source: "
                url = line[len(prefix):]
                para.add_run(prefix)
                r_id = result_doc.part.relate_to(
                    url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
                )
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("r:id"), r_id)
                run_elem = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                color = OxmlElement("w:color")
                color.set(qn("w:val"), "0000FF")
                rPr.append(color)
                u = OxmlElement("w:u")
                u.set(qn("w:val"), "single")
                rPr.append(u)
                run_elem.append(rPr)
                text_elem = OxmlElement("w:t")
                text_elem.text = url
                run_elem.append(text_elem)
                hyperlink.append(run_elem)
                para._p.append(hyperlink)
                continue

            if line.startswith("Search Term(s) ==> "):
                prefix = "Search Term(s) ==> "
                rest = line[len(prefix):]
                match = re.match(r"(.+?)( \(match: [A-Z+]+\))$", rest)
                if match:
                    terms_str, mode_str = match.group(1), match.group(2)
                    para.add_run(prefix)
                    run = para.add_run(terms_str)
                    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    para.add_run(mode_str)
                else:
                    para.add_run(line)
                continue

            if line in ("Header:", "Results:"):
                run = para.add_run(line)
                run.bold = True
                in_results = (line == "Results:")
                continue

            is_doc_line = line.startswith("Document:")
            if in_results and not is_doc_line and not line.startswith("(") and not line.startswith("Files WITHOUT"):
                # Match line in results section — apply highlighting
                _add_highlighted_line(para, line)
            elif is_doc_line:
                _add_highlighted_line(para, line, bold=True)
            else:
                para.add_run(line)

    result_doc.save(docx_path)
    return result_doc


def write_pii_scan_report(docx_path, scan_results, folder, elapsed, files_searched):
    """Generate a .docx report from PII scan results with yellow-highlighted matches.

    Each category gets a section with a severity header and per-file match details.
    Matched text is highlighted in yellow.
    """
    from docsearch.sensitive_patterns import SEVERITY_COLORS, SEVERITY_ORDER

    if os.path.exists(docx_path):
        os.remove(docx_path)

    doc = Document()

    # Title
    title = doc.add_heading("Sensitive Data Scan Report", level=1)

    # Summary
    total = sum(r["match_count"] for r in scan_results)
    high = sum(r["match_count"] for r in scan_results if r["severity"] == "high")
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    doc.add_paragraph(f"Folder: {folder}")
    doc.add_paragraph(f"Date: {now}")
    doc.add_paragraph(f"Files scanned: {files_searched}")
    doc.add_paragraph(f"Scan time: {elapsed:.1f}s")
    doc.add_paragraph(f"Total findings: {total}  ({high} high severity)")

    if total == 0:
        para = doc.add_paragraph()
        run = para.add_run("No sensitive data found.")
        run.font.color.rgb = RGBColor(0, 128, 0)
        run.bold = True
        doc.save(docx_path)
        return doc

    # Summary table
    doc.add_heading("Summary", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Severity"
    hdr[1].text = "Category"
    hdr[2].text = "Matches"
    hdr[3].text = "Files"

    severity_rank = {s: i for i, s in enumerate(SEVERITY_ORDER)}
    sorted_results = sorted(
        scan_results,
        key=lambda r: (severity_rank.get(r["severity"], 99), -r["match_count"]),
    )

    for result in sorted_results:
        row = table.add_row().cells
        sev = SEVERITY_COLORS.get(result["severity"], SEVERITY_COLORS["info"])
        row[0].text = sev["label"]
        row[1].text = result["category"]
        row[2].text = str(result["match_count"])
        row[3].text = str(result["file_count"])

        # Color the severity cell
        if result["severity"] == "high" and result["match_count"] > 0:
            for paragraph in row[0].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(204, 0, 0)
                    run.bold = True

    # Detail sections per category with matches
    doc.add_heading("Details", level=2)

    for result in sorted_results:
        if result["match_count"] == 0:
            continue

        sev = SEVERITY_COLORS.get(result["severity"], SEVERITY_COLORS["info"])
        doc.add_heading(f"[{sev['label']}] {result['category']}", level=3)
        doc.add_paragraph(f"{result['description']} — {result['match_count']} match(es) in {result['file_count']} file(s)")

        # Build highlight pattern from this category's regex
        regex = result.get("regex", "")
        combined_pattern = None
        if regex:
            try:
                combined_pattern = re.compile(regex, re.IGNORECASE)
            except re.error:
                pass

        for fname, info in sorted(result["files"].items()):
            line_nums = sorted(set(info["lines"]))[:20]
            lines_str = ", ".join(str(ln) for ln in line_nums)
            if len(info["lines"]) > 20:
                lines_str += "..."

            para = doc.add_paragraph()
            run = para.add_run(f"{fname}")
            run.bold = True
            para.add_run(f"  ({info['count']} match(es) — lines {lines_str})")

            # Add matched text with highlighting
            for match_text in info.get("match_texts", [])[:50]:
                para = doc.add_paragraph()
                if combined_pattern:
                    parts = combined_pattern.split(match_text)
                    matches = combined_pattern.findall(match_text)
                    for i, part in enumerate(parts):
                        if part:
                            para.add_run(part)
                        if i < len(matches):
                            run = para.add_run(matches[i])
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                else:
                    para.add_run(match_text)

    # Disclaimer
    doc.add_paragraph("")
    para = doc.add_paragraph()
    run = para.add_run(
        "This report was generated by docsearch's Sensitive Data Scan feature. "
        "Pattern-based detection may produce false positives (e.g., a 9-digit number "
        "that is not an SSN). Review each finding to determine whether it represents "
        "actual sensitive data."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(docx_path)
    return doc


def insert_file_sizes(txt_path, docx_path, result_doc):
    """Insert report file sizes into both txt and docx reports.

    Returns (txt_size, docx_size) in bytes.
    """
    txt_size = os.path.getsize(txt_path)
    docx_size = os.path.getsize(docx_path)
    txt_name = os.path.basename(txt_path)
    docx_name = os.path.basename(docx_path)
    sizes_line = f"Report File Sizes ==> {txt_name} ({fmt_size(txt_size)}), {docx_name} ({fmt_size(docx_size)})"

    # Update txt report
    with open(txt_path, "r", encoding="utf-8") as f:
        content = f.read()
    content = content.replace(
        "\nCommand ==>",
        f"\n{sizes_line}\nCommand ==>",
        1,
    )
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(content)

    # Update docx report — insert sizes paragraph after timestamp
    for i, para in enumerate(result_doc.paragraphs):
        if para.text.startswith("Report Generated On ==>"):
            new_para = OxmlElement("w:p")
            run_elem = OxmlElement("w:r")
            text_elem = OxmlElement("w:t")
            text_elem.text = sizes_line
            run_elem.append(text_elem)
            new_para.append(run_elem)
            para._p.addnext(new_para)
            break
    result_doc.save(docx_path)

    return (txt_size, docx_size)


def write_csv_report(output_path, matches, inverse_files=None):
    """Write docsearch_results.csv."""
    if os.path.exists(output_path):
        os.remove(output_path)
    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        if inverse_files is not None:
            writer.writerow(["filename", "folder"])
            for fp in inverse_files:
                writer.writerow([os.path.basename(fp), os.path.dirname(fp)])
        else:
            writer.writerow(["filename", "folder", "line_number", "matched_text"])
            for file_dir, filename, line_num, text in matches:
                writer.writerow([filename, file_dir, line_num, _strip_highlights(text)])


def write_json_report(output_path, matches, search_terms, report_mode,
                      files_count, search_elapsed, inverse_files=None):
    """Write docsearch_results.json."""
    if os.path.exists(output_path):
        os.remove(output_path)

    if inverse_files is not None:
        json_data = {
            "search_terms": search_terms,
            "mode": report_mode,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "files_searched": files_count,
            "files_without_matches": len(inverse_files),
            "elapsed_seconds": round(search_elapsed, 2),
            "inverse_files": [
                {
                    "filename": os.path.basename(fp),
                    "folder": os.path.dirname(fp),
                }
                for fp in inverse_files
            ],
        }
    else:
        # Per-file match counts (ordered by first appearance)
        file_counts = {}
        for file_dir, filename, _ln, _tx in matches:
            key = (file_dir, filename)
            if key not in file_counts:
                file_counts[key] = 0
            file_counts[key] += 1
        matches_per_file = [
            {"filename": fn, "folder": fd, "matches": count}
            for (fd, fn), count in file_counts.items()
        ]

        json_data = {
            "search_terms": search_terms,
            "mode": report_mode,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "files_searched": files_count,
            "matches_found": len(matches),
            "matches_per_file": matches_per_file,
            "elapsed_seconds": round(search_elapsed, 2),
            "matches": [
                {
                    "filename": filename,
                    "folder": file_dir,
                    "line_number": line_num,
                    "matched_text": _strip_highlights(text),
                }
                for file_dir, filename, line_num, text in matches
            ],
        }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(json_data, f, indent=2, ensure_ascii=False)


def _pdf_safe(text):
    """Replace Unicode characters that Helvetica can't render."""
    replacements = {
        '\u2018': "'", '\u2019': "'",   # smart single quotes
        '\u201c': '"', '\u201d': '"',   # smart double quotes
        '\u2013': '-', '\u2014': '--',  # en dash, em dash
        '\u2026': '...', '\u00a0': ' ', # ellipsis, nbsp
        '\u2022': '*', '\u2023': '>',   # bullets
        '\ufb01': 'fi', '\ufb02': 'fl', # ligatures
        '\u00b0': 'deg',               # degree sign
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Remove any remaining non-latin1 characters
    return text.encode('latin-1', errors='replace').decode('latin-1')


def _pdf_insert_highlighted(pdf, text, search_terms):
    """Insert text with search terms highlighted in orange background."""
    import re as _re
    if not search_terms:
        pdf.multi_cell(0, 5, text)
        return
    # Build pattern from search terms
    patterns = []
    for term in search_terms:
        patterns.append(_re.escape(term))
    try:
        pattern = _re.compile("|".join(patterns), _re.IGNORECASE)
    except _re.error:
        pdf.multi_cell(0, 5, text)
        return
    parts = pattern.split(text)
    found = pattern.findall(text)
    if not found:
        pdf.multi_cell(0, 5, text)
        return
    # Use cell for inline highlighting (multi_cell doesn't support mixed colors)
    x_start = pdf.get_x()
    page_width = pdf.w - pdf.r_margin - x_start
    for i, part in enumerate(parts):
        if part:
            pdf.set_text_color(0, 0, 0)
            pdf.write(5, part)
        if i < len(found):
            # Yellow background highlight
            pdf.set_fill_color(255, 255, 0)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(pdf.get_string_width(found[i]) + 1, 5, found[i], fill=True)
            pdf.set_fill_color(255, 255, 255)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(5)


def write_pdf_report(output_path, matches, search_terms=None,
                     report_mode="ANY", inverse_files=None):
    """Write docsearch_results.pdf with highlighted matches."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "docsearch Search Results", ln=True)
    pdf.ln(3)

    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, _pdf_safe(f"Search terms: {' '.join(search_terms or [])}"), ln=True)
    pdf.cell(0, 6, f"Mode: {report_mode}", ln=True)
    pdf.cell(0, 6, f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
    pdf.ln(5)

    if inverse_files is not None:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, f"Files WITHOUT matches: {len(inverse_files)}", ln=True)
        pdf.ln(3)
        pdf.set_font("Helvetica", "", 10)
        for fp in inverse_files:
            filename = os.path.basename(fp)
            dirname = os.path.dirname(fp)
            pdf.cell(0, 5, _pdf_safe(f"  {filename}"), ln=True)
            pdf.set_text_color(128, 128, 128)
            pdf.cell(0, 5, _pdf_safe(f"  ({dirname})"), ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(2)
    else:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, f"Matches found: {len(matches)}", ln=True)
        pdf.ln(3)

        prev_file = None
        for file_dir, filename, line_num, text in matches:
            current_file = os.path.join(file_dir, filename)
            if current_file != prev_file:
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 10)
                fc = sum(1 for fd, fn, _ln, _tx in matches
                         if os.path.join(fd, fn) == current_file)
                pdf.set_text_color(26, 115, 232)
                pdf.cell(0, 6, _pdf_safe(f"Document: {filename} ({fc} match{'es' if fc != 1 else ''})"), ln=True)
                pdf.set_text_color(128, 128, 128)
                pdf.set_font("Helvetica", "", 9)
                pdf.cell(0, 5, _pdf_safe(f"({file_dir})"), ln=True)
                pdf.set_text_color(0, 0, 0)
                prev_file = current_file

            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(128, 128, 128)
            pdf.cell(15, 5, f"Line {line_num}:")
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 10)
            clean_text = _pdf_safe(_strip_highlights(text))
            if len(clean_text) > 200:
                clean_text = clean_text[:197] + "..."
            _pdf_insert_highlighted(pdf, clean_text, [_pdf_safe(t) for t in (search_terms or [])])
            pdf.ln(1)

    pdf.output(output_path)


def append_results(append_name, output_dir, txt_path, docx_path):
    """Append current results to accumulated DO_NOT_SEARCH files."""
    append_txt_path = os.path.join(output_dir, f"DO_NOT_SEARCH_ACCUMULATED_{append_name}.txt")
    append_docx_path = os.path.join(output_dir, f"DO_NOT_SEARCH_ACCUMULATED_{append_name}.docx")
    with open(txt_path, "r", encoding="utf-8") as src:
        results_content = src.read()
    with open(append_txt_path, "a", encoding="utf-8") as dst:
        dst.write(results_content)
    if os.path.exists(append_docx_path):
        existing_doc = Document(append_docx_path)
        new_doc = Document(docx_path)
        body = existing_doc.element.body
        sect_pr = body.find(qn('w:sectPr'))
        for para in new_doc.paragraphs:
            new_elem = deepcopy(para._p)
            if sect_pr is not None:
                sect_pr.addprevious(new_elem)
            else:
                body.append(new_elem)
        existing_doc.save(append_docx_path)
    else:
        shutil.copy2(docx_path, append_docx_path)


# ── Suite Reports ──────────────────────────────────────────────

def _describe_search(result):
    """Return a short description of a search result's configuration."""
    parts = [f'"{result["search_text"]}"']
    if result.get("inverse"):
        parts.append("(inverse)")
    return " ".join(parts)


def write_suite_report_txt(output_path, suite_name, folder, results, start_time, end_time):
    """Write a search suite compliance/audit report as a text file."""
    ts = datetime.fromtimestamp(start_time).strftime("%Y-%m-%d %H:%M:%S")
    elapsed = end_time - start_time
    passed = sum(1 for r in results if r["passed"])
    total = len(results)
    verdict = "PASSED" if passed == total else "FAILED"

    lines = [
        "=" * 50,
        "docsearch Search Suite Report",
        "=" * 50,
        f"Suite: {suite_name}",
        f"Folder: {folder}",
        f"Date: {ts}",
        f"Duration: {elapsed:.1f} seconds",
    ]
    has_cascade = any(r.get("cascade_file_count") is not None for r in results)
    if has_cascade:
        lines.append("Mode: Cascade (each stage narrows files from previous)")
    lines += [
        "",
        "Results:",
        "-" * 40,
    ]

    for r in results:
        status = "PASS" if r["passed"] else "FAIL"
        lines.append(f"  [{status}] {r['name']}")
        lines.append(f"    Search: {_describe_search(r)}")
        if r["passed"]:
            if r.get("inverse"):
                lines.append(f"    Result: {r['match_count']} file(s) without matches")
            else:
                result_str = f"    Result: {r['match_count']} match(es) found"
                cfc = r.get("cascade_file_count")
                cic = r.get("cascade_input_count")
                if cfc is not None:
                    result_str += f" in {cfc} file(s)"
                    if cic is not None:
                        result_str += f" (narrowed from {cic})"
                lines.append(result_str)
        elif r.get("return_code") == 2:
            lines.append(f"    Result: Error — {r.get('summary', 'unknown error')}")
        elif r.get("inverse"):
            lines.append("    Result: All files matched (none missing)")
        else:
            lines.append("    Result: No matches found")
        stage_files = r.get("stage_files", {})
        if stage_files:
            fnames = ", ".join(os.path.basename(p) for p in sorted(stage_files.values()))
            lines.append(f"    Stage reports: {fnames}")
        lines.append("")

    lines.append("=" * 50)
    lines.append(f"Summary: {passed} of {total} tests passed. {verdict}")
    lines.append("=" * 50)
    lines.append("")
    lines.append("Disclaimer: This report was generated by docsearch, a search and")
    lines.append("reporting tool. It does not constitute legal, regulatory, or compliance")
    lines.append("advice. Pass/fail results indicate whether search criteria were met, not")
    lines.append("whether documents satisfy regulatory requirements.")
    lines.append("")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def write_suite_report_json(output_path, suite_name, folder, results, start_time, end_time):
    """Write suite results as JSON for programmatic consumption."""
    ts = datetime.fromtimestamp(start_time).strftime("%Y-%m-%dT%H:%M:%S")
    elapsed = end_time - start_time
    passed = sum(1 for r in results if r["passed"])
    total = len(results)
    verdict = "PASSED" if passed == total else "FAILED"

    has_cascade = any(r.get("cascade_file_count") is not None for r in results)
    data = {
        "suite_name": suite_name,
        "folder": folder,
        "timestamp": ts,
        "duration_seconds": round(elapsed, 2),
        "total_tests": total,
        "passed": passed,
        "failed": total - passed,
        "overall": verdict,
        "cascade": has_cascade,
        "tests": [
            {
                "name": r["name"],
                "search_text": r["search_text"],
                "inverse": r.get("inverse", False),
                "passed": r["passed"],
                "match_count": r["match_count"],
                "return_code": r["return_code"],
                "stage_files": {
                    ext: os.path.basename(path)
                    for ext, path in r.get("stage_files", {}).items()
                },
                "cascade_input_count": r.get("cascade_input_count"),
                "cascade_file_count": r.get("cascade_file_count"),
            }
            for r in results
        ],
    }

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def write_suite_report_docx(output_path, suite_name, folder, results,
                            start_time, end_time, version=None,
                            source_files=None):
    """Write a consolidated suite report as a formatted .docx file.

    Produces an evidence-grade document with suite metadata, a color-coded
    summary table (green PASS / red FAIL), per-stage details, a file-set
    fingerprint for audit traceability, and a source file manifest.

    source_files: optional list of (filepath, size_bytes, modified_time_str)
                  tuples for the documents that were searched.
    """
    import hashlib

    ts = datetime.fromtimestamp(start_time).strftime("%Y-%m-%d %H:%M:%S")
    elapsed = end_time - start_time
    passed = sum(1 for r in results if r["passed"])
    total = len(results)
    verdict = "PASSED" if passed == total else "FAILED"
    has_cascade = any(r.get("cascade_file_count") is not None for r in results)

    doc = Document()

    # ── Styles ──
    style = doc.styles["Normal"]
    style.font.size = Pt(14)
    style.font.name = "Calibri"

    GREEN = RGBColor(0, 128, 0)
    RED = RGBColor(200, 0, 0)

    def _add_line(text, bold=False, size=None, color=None):
        para = doc.add_paragraph()
        run = para.add_run(text)
        if bold:
            run.bold = True
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        return para

    # ── Title ──
    _add_line("docsearch Search Suite Report", bold=True, size=16)
    doc.add_paragraph()  # spacer

    # ── Metadata ──
    _add_line(f"Suite:     {suite_name}", bold=True)
    _add_line(f"Folder:    {folder}")
    _add_line(f"Date:      {ts}")
    _add_line(f"Duration:  {elapsed:.1f} seconds")
    if has_cascade:
        _add_line("Mode:      Cascade (each stage narrows files from previous)")
    if version:
        _add_line(f"Version:   docsearch {version}")

    # ── File-set fingerprint ──
    # Hash sorted (basename, size) of all stage report source files for traceability.
    file_entries = []
    for r in results:
        for path in r.get("stage_files", {}).values():
            if os.path.exists(path):
                file_entries.append((os.path.basename(path), os.path.getsize(path)))
    if file_entries:
        file_entries.sort()
        digest = hashlib.sha256(
            "".join(f"{n}:{s}" for n, s in file_entries).encode()
        ).hexdigest()[:16]
        _add_line(f"Report fingerprint: {digest}")

    doc.add_paragraph()  # spacer

    # ── Verdict banner ──
    verdict_color = GREEN if verdict == "PASSED" else RED
    para = doc.add_paragraph()
    run = para.add_run(f"Result: {passed} of {total} tests passed — {verdict}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = verdict_color

    doc.add_paragraph()  # spacer

    # ── Source file manifest (moved above summary) ──
    if source_files:
        _add_line("Source File Manifest", bold=True, size=12)
        _add_line(
            f"{len(source_files)} file(s) were present in the search folder "
            f"when this suite ran."
        )
        doc.add_paragraph()

        manifest_table = doc.add_table(rows=1, cols=4)
        manifest_table.style = "Light Grid Accent 1"
        for i, h in enumerate(["#", "File", "Size", "Last Modified"]):
            cell = manifest_table.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)

        for idx, (filepath, size_bytes, mod_time) in enumerate(source_files, 1):
            row = manifest_table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = os.path.basename(filepath)
            row.cells[2].text = fmt_size(size_bytes)
            row.cells[3].text = mod_time
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.size is None:
                            run.font.size = Pt(10)

        for row in manifest_table.rows:
            trPr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            trPr.append(cant_split)

        total_size = sum(s for _, s, _ in source_files)
        doc.add_paragraph()
        _add_line(f"Total: {len(source_files)} files, {fmt_size(total_size)}")
        doc.add_paragraph()

    # ── Summary table ──
    _add_line("Test Summary", bold=True, size=12)

    cols = 5 if has_cascade else 4
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Light Grid Accent 1"

    headers = ["#", "Test", "Status", "Matches"]
    if has_cascade:
        headers.append("Files")
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    for idx, r in enumerate(results, 1):
        row = table.add_row()
        row.cells[0].text = str(idx)

        row.cells[1].text = r["name"]

        status = "PASS" if r["passed"] else "FAIL"
        status_cell = row.cells[2]
        status_cell.text = ""
        run = status_cell.paragraphs[0].add_run(status)
        run.bold = True
        run.font.color.rgb = GREEN if r["passed"] else RED
        run.font.size = Pt(9)

        if r.get("inverse"):
            row.cells[3].text = f"{r['match_count']} file(s)"
        else:
            row.cells[3].text = str(r["match_count"])

        if has_cascade:
            cfc = r.get("cascade_file_count")
            cic = r.get("cascade_input_count")
            if cfc is not None:
                cascade_str = str(cfc)
                if cic is not None:
                    cascade_str += f" (of {cic})"
                row.cells[4].text = cascade_str

        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.size is None:
                        run.font.size = Pt(9)

    doc.add_paragraph()  # spacer

    # ── File × Test matrix ──
    if source_files and results:
        _add_line("File × Test Detail", bold=True, size=12)
        _add_line(
            "Each cell shows the match count for that file in that test. "
            "For inverse tests, \u2718 = file is missing the required term. "
            "The last column shows PASS if all tests passed for that file, "
            "or FAIL if any test failed."
        )
        doc.add_paragraph()

        # Build per-file, per-test match data
        test_names = [r["name"] for r in results]
        # Map: test_name -> {filename -> count}
        test_file_counts = {}
        test_is_inverse = {}
        for r in results:
            tname = r["name"]
            test_is_inverse[tname] = r.get("inverse", False)
            file_map = {}
            for item in r.get("matched_files", []):
                fname = item[0] if isinstance(item, (list, tuple)) else item
                count = item[1] if isinstance(item, (list, tuple)) and len(item) > 1 else 0
                file_map[fname] = count
            test_file_counts[tname] = file_map

        # Collect all source filenames
        all_filenames = [os.path.basename(fp) for fp, _, _ in source_files]

        # Build matrix table: #, File, Test1, Test2, ..., P/F
        n_tests = len(test_names)
        matrix = doc.add_table(rows=1, cols=2 + n_tests + 1)
        matrix.style = "Light Grid Accent 1"

        # Row 0: header with test numbers
        header_row = matrix.rows[0]
        cell = header_row.cells[0]
        cell.text = ""
        run = cell.paragraphs[0].add_run("#")
        run.bold = True
        run.font.size = Pt(10)
        cell = header_row.cells[1]
        cell.text = ""
        run = cell.paragraphs[0].add_run("Test \u2192")
        run.bold = True
        run.font.size = Pt(10)
        for t_idx in range(n_tests):
            cell = header_row.cells[2 + t_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(t_idx + 1))
            run.bold = True
            run.font.size = Pt(10)
        cell = header_row.cells[2 + n_tests]
        cell.text = ""
        run = cell.paragraphs[0].add_run("P/F")
        run.bold = True
        run.font.size = Pt(10)

        for file_idx, fname in enumerate(all_filenames, 1):
            row = matrix.add_row()
            row.cells[0].text = str(file_idx)
            row.cells[1].text = fname

            file_failed = False
            for t_idx, tname in enumerate(test_names):
                cell = row.cells[2 + t_idx]
                file_map = test_file_counts.get(tname, {})
                is_inv = test_is_inverse.get(tname, False)
                pc = results[t_idx].get("pass_criteria", {"op": ">=", "n": 1})

                if is_inv:
                    # Inverse: file in matched_files means it's MISSING the term
                    if fname in file_map:
                        cell.text = ""
                        run = cell.paragraphs[0].add_run("\u2718")
                        run.font.color.rgb = RED
                        run.font.size = Pt(10)
                        # Missing term — check if criteria expects 0 missing
                        if pc["op"] == "==" and pc["n"] == 0:
                            file_failed = True
                    else:
                        cell.text = ""
                        run = cell.paragraphs[0].add_run("\u2714")
                        run.font.color.rgb = GREEN
                        run.font.size = Pt(10)
                else:
                    # Normal: file in matched_files has matches
                    count = file_map.get(fname, 0)
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(str(count))
                    run.font.size = Pt(10)
                    if count > 0:
                        # Has matches — fail if criteria wants 0
                        if pc["op"] == "==" and pc["n"] == 0:
                            run.font.color.rgb = RED
                            file_failed = True
                        elif pc["op"] == "<=" and count > pc["n"]:
                            run.font.color.rgb = RED
                            file_failed = True
                    else:
                        # No matches — fail if criteria requires matches
                        if pc["op"] == ">=" and pc["n"] > 0:
                            run.font.color.rgb = RED
                            file_failed = True

            # P/F column
            pf_cell = row.cells[2 + n_tests]
            pf_cell.text = ""
            if file_failed:
                run = pf_cell.paragraphs[0].add_run("F")
                run.bold = True
                run.font.color.rgb = RED
            else:
                run = pf_cell.paragraphs[0].add_run("P")
                run.bold = True
                run.font.color.rgb = GREEN
            run.font.size = Pt(10)

            # Set font size on # and File cells
            for ci in (0, 1):
                for para in row.cells[ci].paragraphs:
                    for run in para.runs:
                        if run.font.size is None:
                            run.font.size = Pt(10)

        # Summary row: P/F status per test column
        summary_row = matrix.add_row()
        summary_row.cells[0].text = ""
        cell = summary_row.cells[1]
        cell.text = ""
        run = cell.paragraphs[0].add_run("Test Result:")
        run.bold = True
        run.font.size = Pt(10)
        for t_idx, r in enumerate(results):
            cell = summary_row.cells[2 + t_idx]
            cell.text = ""
            status = "P" if r["passed"] else "F"
            run = cell.paragraphs[0].add_run(status)
            run.bold = True
            run.font.color.rgb = GREEN if r["passed"] else RED
            run.font.size = Pt(10)
        summary_row.cells[2 + n_tests].text = ""

        # Prevent rows from splitting across pages; keep last data row
        # attached to the summary row
        all_rows = list(matrix.rows)
        for i, row in enumerate(all_rows):
            trPr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            trPr.append(cant_split)
            # keepNext on second-to-last row keeps summary row attached
            if i == len(all_rows) - 2:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        pPr = para._p.get_or_add_pPr()
                        keep_next = OxmlElement("w:keepNext")
                        pPr.append(keep_next)

        doc.add_paragraph()

    # ── Stage details ──
    _add_line("Stage Details", bold=True, size=12)

    for idx, r in enumerate(results, 1):
        status = "PASS" if r["passed"] else "FAIL"
        status_color = GREEN if r["passed"] else RED

        para = doc.add_paragraph()
        run = para.add_run(f"Stage {idx}: {r['name']}  —  ")
        run.bold = True
        run.font.size = Pt(10)
        run = para.add_run(status)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = status_color

        # Search description
        search_desc = f"Search: {_describe_search(r)}"
        pc = r.get("pass_criteria", {})
        if pc:
            search_desc += f"  [criteria: {pc.get('op', '>=')} {pc.get('n', 1)}]"
        _add_line(search_desc)

        # Result
        if r.get("return_code") == 2:
            _add_line(f"Result: Error — {r.get('summary', 'unknown error')}", color=RED)
        elif r.get("inverse"):
            _add_line(f"Result: {r['match_count']} file(s) without matches")
        else:
            result_str = f"Result: {r['match_count']} match(es) found"
            cfc = r.get("cascade_file_count")
            cic = r.get("cascade_input_count")
            if cfc is not None:
                result_str += f" in {cfc} file(s)"
                if cic is not None:
                    result_str += f" (narrowed from {cic})"
            _add_line(result_str)

        # Stage report references
        stage_files = r.get("stage_files", {})
        if stage_files:
            fnames = ", ".join(os.path.basename(p) for p in sorted(stage_files.values()))
            _add_line(f"Stage reports: {fnames}")

        doc.add_paragraph()  # spacer between stages

    # ── Disclaimer ──
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    run = disclaimer.add_run(
        "Disclaimer: This report was generated by docsearch, a search and reporting tool. "
        "It does not constitute legal, regulatory, or compliance advice. Pass/fail results "
        "indicate whether search criteria were met, not whether documents satisfy regulatory "
        "requirements. Users are solely responsible for determining whether results meet "
        "their specific compliance obligations."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(output_path)
