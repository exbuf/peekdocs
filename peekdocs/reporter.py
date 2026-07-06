"""Report generation for peekdocs (TXT, DOCX, CSV, JSON, PDF, append)."""
from __future__ import annotations

import csv
import json
import os
import re
import shutil
import stat
import sys
import textwrap
from copy import deepcopy
from datetime import datetime
from typing import Any


# Match "tuple" used throughout the reporter: (file_dir, filename,
# line_num, text). Accepts both plain 4-tuples (from cli.py) and
# SearchMatch dataclasses (from api.py). SearchMatch defines __iter__ +
# __getitem__ + __len__ so it unpacks and indexes identically to a plain
# tuple at runtime, but mypy sees them as distinct types. The reporter is
# a polymorphic consumer — Any is honest about that.
Match = Any

# Module-level flag — when True, report files are set to owner-only
# read/write (chmod 600) after creation.  The GUI sets this based on
# the "Restrict file permissions" checkbox.
restrict_permissions: bool = False


def _restrict_file_permissions(filepath: str) -> None:
    """Set owner-only read/write permissions on a file (Unix/macOS).

    Only applies when ``restrict_permissions`` is True.  On Windows
    this is always a no-op — NTFS permissions are managed differently.
    """
    if not restrict_permissions:
        return
    if sys.platform != "win32":
        try:
            os.chmod(filepath, stat.S_IRUSR | stat.S_IWUSR)
        except OSError:
            pass

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor, Inches

from peekdocs.scanner import _wildcard_to_regex
from peekdocs.translator import translate_search


def fmt_size(b: int) -> str:
    """Format byte count as human-readable string."""
    if b >= 1_000_000:
        return f"{b / 1_000_000:.2f} MB"
    elif b >= 1_000:
        return f"{b / 1_000:.2f} KB"
    return f"{b} bytes"


def _strip_highlights(text: str) -> str:
    """Remove ** highlight markers from text."""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _sha256_of_file(path: str) -> str | None:
    """Return the SHA-256 hex digest of *path*'s raw bytes, or None on read error.

    Reads the file in 64 KiB chunks so multi-GB files don't load into memory.
    Used by --hash to add a content fingerprint to each matched-file entry in
    the JSON output. Any OS-level read failure (missing, permission denied,
    deleted between search and hash) yields None so the search itself isn't
    aborted.
    """
    import hashlib
    try:
        h = hashlib.sha256()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(65536), b""):
                h.update(chunk)
        return h.hexdigest()
    except OSError:
        return None


def write_txt_report(
    output_path: str,
    matches: list[Match],
    all_files: list[str],
    search_terms: list[str],
    command_str: str,
    report_mode: str,
    use_ocr: bool,
    exclude_terms: list[str] | None,
    use_context: bool,
    use_fuzzy: bool,
    use_regex: bool,
    use_wildcard: bool,
    search_elapsed: float,
    cores: int,
    cpu_count: int,
    inverse_files: list[str] | None = None,
    recursive: bool = False,
    file_types: list[str] | str | None = None,
    proximity: int | None = None,
    context_before: int = 0,
    context_after: int = 0,
    specific_files: list[str] | str | None = None,
    use_index: bool = False,
    inverse: bool = False,
    output_csv: bool = False,
    output_json: bool = False,
    expression: str | None = None,
    use_whole_word: bool = False,
    total_matches: int | None = None,
    max_matches: int | None = None,
    range_specs: list[Any] | None = None,
    index_meta: dict[str, Any] | None = None,
    bulleted_terms: bool = False,
    pattern_sections: list[dict[str, Any]] | None = None,
) -> tuple[int, str]:
    """Write the .txt result report (e.g. peekdocs_standard_results.txt).

    Returns (total_bytes, size_str) for use in console summary.
    """
    if os.path.exists(output_path):
        os.remove(output_path)
    with open(output_path, "w", encoding="utf-8") as f:
        from peekdocs.cli import VERSION as _ver
        f.write(f"peekdocs v{_ver}\n\n\n")
        f.write(f"Report Generated On ==> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Saved as ==> {os.path.abspath(output_path)}\n")
        f.write("NOTE: This report is overwritten after each new search. To keep a permanent copy, use 'Save report as:' in Advanced Search Options (GUI) or the -s flag (CLI).\n")
        f.write(f"Command ==> {command_str}\n")
        translation = translate_search(
            search_terms, report_mode=report_mode,
            use_regex=use_regex, use_wildcard=use_wildcard,
            use_fuzzy=use_fuzzy, use_ocr=use_ocr, use_index=use_index,
            use_whole_word=use_whole_word,
            inverse=inverse, recursive=recursive,
            exclude_terms=exclude_terms,
            file_types=file_types, specific_files=specific_files,
            proximity=proximity,
            context_before=context_before, context_after=context_after,
            expression=expression,
            range_specs=range_specs,
        )
        f.write(f"Translation ==> {translation}\n")
        if expression:
            f.write(f"Search Expression ==> {expression} (mode: {report_mode})\n")
        elif bulleted_terms and len(search_terms) > 1:
            # Regex collections / multi-term searches: bullet each term on
            # its own line for readability. The flat space-joined form
            # collapses 10+ regex patterns into one wall of characters
            # that's nearly impossible to scan.
            f.write(f"Search Term(s) (match: {report_mode}):\n")
            for _t in search_terms:
                f.write(f"  • {_t}\n")
        else:
            f.write(f"Search Term(s) ==> {' '.join(search_terms)} (match: {report_mode})\n")
        if exclude_terms:
            f.write(f"Exclude Term(s) ==> {' '.join(exclude_terms)}\n")
        if pattern_sections is not None:
            # Per-pattern render: Hits reflects the true sum across all
            # pattern sections, not the (possibly deduped/capped) flat
            # `matches` list. Every match shown below.
            _hits = sum(s.get("match_count", 0) for s in pattern_sections)
            f.write(f"Hits ==> {_hits} (across {len(pattern_sections)} pattern(s) — see per-pattern sections below)\n")
        elif total_matches is not None and total_matches > len(matches):
            f.write(f"Hits ==> {len(matches)} (of {total_matches:,} total — report capped at {max_matches:,}; use -m to change)\n")
        else:
            f.write(f"Hits ==> {len(matches)}\n")
        if len(matches) == 0 and not inverse:
            f.write(f"\n*** NO MATCHES FOUND ***\n")
            f.write(f"Searched {len(all_files)} file(s). Check your search terms, folder, and settings above.\n")

        # ── Search Settings ──
        on_off = lambda v: "ON" if v else "OFF"
        f.write(f"\nSearch Settings:\n")
        f.write(f"  AND mode: {on_off(report_mode == 'ALL')}  |  Recursive: {on_off(recursive)}  |  Inverse: {on_off(inverse)}  |  Expression: {on_off(expression is not None)}\n")
        f.write(f"  Fuzzy: {on_off(use_fuzzy)}  |  Wildcard: {on_off(use_wildcard)}  |  Regex: {on_off(use_regex)}  |  Whole Word: {on_off(use_whole_word)}  |  OCR: {on_off(use_ocr)}\n")
        f.write(f"  Index: {on_off(use_index)}\n")
        if use_index and index_meta:
            f.write(f"  Index last updated: {index_meta.get('last_updated', index_meta.get('created_at', '?'))}"
                    f"  ({index_meta.get('file_count', '?')} files, {index_meta.get('line_count', '?')} lines)\n")
        if file_types:
            f.write(f"  File types: {file_types}\n")
        if specific_files:
            f.write(f"  Specific files: {specific_files}\n")
        if proximity:
            f.write(f"  Proximity: {proximity} words\n")
        if use_context:
            f.write(f"  Context: {context_before} line(s) before, {context_after} line(s) after\n")
        outputs = ["TXT", "DOCX"]
        if output_csv:
            outputs.append("CSV")
        if output_json:
            outputs.append("JSON")
        f.write(f"  Output formats: {', '.join(outputs)}\n\n")
        # Per-file match counts (used later for per-match headers)
        file_counts = {}
        for fd, fn, _ln, _tx in matches:
            key = (fd, fn)
            if key not in file_counts:
                file_counts[key] = 0
            file_counts[key] += 1
        f.write(f"Search Time ==> {search_elapsed:.2f} seconds, Cores used ==> {cores} of {cpu_count}\n")
        total_bytes = sum(os.path.getsize(f_path) for f_path in all_files)
        if total_bytes >= 1_000_000:
            size_str = f"{total_bytes / 1_000_000:.2f} MB"
        elif total_bytes >= 1_000:
            size_str = f"{total_bytes / 1_000:.2f} KB"
        else:
            size_str = f"{total_bytes} bytes"
        f.write(f"Files searched ==> {len(all_files)} ({size_str})\n")
        ext_counts: dict[str, int] = {}
        for fp in all_files:
            ext = os.path.splitext(fp)[1].lower()
            ext_counts[ext] = ext_counts.get(ext, 0) + 1
        if ext_counts:
            tally = ", ".join(f"{ext}: {count}" for ext, count in ext_counts.items())
            f.write(f"File Types Searched ==> {tally}\n")
        f.write("\n")
        f.write("Results:\n\n")
        if inverse_files is not None:
            f.write(f"Files WITHOUT matches: {len(inverse_files)} (out of {len(all_files)} searched)\n\n")
            for fp in inverse_files:
                f.write(f"  {os.path.basename(fp)}\n")
                f.write(f"  ({os.path.dirname(fp)})\n\n")
            # Inverse reports only list files missing the term — no match lines
            return (total_bytes, size_str)
        if pattern_sections is not None:
            # Per-pattern render path (regex collections / multi-pattern).
            # Each pattern gets a clearly-headed section listing every
            # one of its matches; file-count labels reset per pattern so
            # 'X matches in Y file(s)' reflects that pattern alone, not
            # the cumulative across earlier patterns.
            for _ps_idx, _section in enumerate(pattern_sections):
                _sec_matches = _section.get("matches") or []
                _sec_name = _section.get("name", f"Pattern {_ps_idx + 1}")
                _sec_regex = _section.get("regex", "")
                _sec_mc = _section.get("match_count", len(_sec_matches))
                _sec_fc = _section.get("file_count", 0)
                # Per-pattern file-count map for the 'Document: X (N
                # matches), Line: ...' label inside this section.
                _sec_file_counts: dict[tuple[str, str], int] = {}
                for _fd, _fn, _ln, _tx in _sec_matches:
                    _key = (_fd, _fn)
                    _sec_file_counts[_key] = _sec_file_counts.get(_key, 0) + 1
                f.write("=" * 72 + "\n")
                f.write(f"Pattern: {_sec_name}\n")
                f.write(f"Regex:   {_sec_regex}\n")
                _files_word = "file" if _sec_fc == 1 else "files"
                _matches_word = "match" if _sec_mc == 1 else "matches"
                f.write(f"Hits:    {_sec_mc} {_matches_word} in {_sec_fc} {_files_word}\n")
                f.write("=" * 72 + "\n\n")
                if not _sec_matches:
                    f.write("(no matches for this pattern)\n\n")
                    continue
                _prev_file = None
                for _fd, _fn, _ln, _tx in _sec_matches:
                    _current_file = os.path.join(_fd, _fn)
                    if use_context and _prev_file == _current_file:
                        f.write("---\n\n")
                    _prev_file = _current_file
                    if use_context or "\n" in _tx:
                        _lines = _tx.split("\n")
                        _wrapped_lines = [textwrap.fill(_l, width=80) if _l else _l for _l in _lines]
                        _wrapped = "\n".join(_wrapped_lines)
                    else:
                        _wrapped = textwrap.fill(_tx, width=80)
                    _fc = _sec_file_counts.get((_fd, _fn), 1)
                    f.write(
                        f'Document: {_fn} ({_fc} match{"es" if _fc != 1 else ""}), Line: {_ln}, Match:\n'
                        f'({_fd})\n"{_wrapped}"\n\n'
                    )
        else:
            prev_file = None
            for file_dir, filename, line_num, text in matches:
                current_file = os.path.join(file_dir, filename)
                if use_context and prev_file == current_file:
                    f.write("---\n\n")
                prev_file = current_file
                if use_context or "\n" in text:
                    lines = text.split("\n")
                    wrapped_lines = [textwrap.fill(line, width=80) if line else line for line in lines]
                    wrapped = "\n".join(wrapped_lines)
                else:
                    wrapped = textwrap.fill(text, width=80)
                fc = file_counts.get((file_dir, filename), 1)
                f.write(f'Document: {filename} ({fc} match{"es" if fc != 1 else ""}), Line: {line_num}, Match:\n({file_dir})\n"{wrapped}"\n\n')

    _restrict_file_permissions(output_path)
    return (total_bytes, size_str)


def write_docx_report(
    docx_path: str,
    txt_path: str,
    search_terms: list[str] | None = None,
    use_regex: bool = False,
    use_wildcard: bool = False,
    use_whole_word: bool = False,
    use_fuzzy: bool = False,
    expression: str | None = None,
) -> Any:
    """Create the .docx result report from the .txt report with yellow highlighting.

    Highlights matched terms directly using the search parameters rather than
    parsing markers from the text. Returns the Document object for further
    modification.
    """
    if os.path.exists(docx_path):
        os.remove(docx_path)

    # Build highlight patterns from search terms
    highlight_patterns = []
    if not use_fuzzy:
        if expression:
            from peekdocs.expr_parser import parse_expression, extract_positive_terms
            highlight_terms = extract_positive_terms(parse_expression(expression))
        elif search_terms:
            highlight_terms = search_terms
        else:
            highlight_terms = []
        for term in highlight_terms:
            if use_wildcard:
                highlight_patterns.append(_wildcard_to_regex(term))
            elif use_regex:
                highlight_patterns.append(term)
            elif use_whole_word:
                prefix = r'\b' if re.match(r'\w', term) else ''
                suffix = r'\b' if re.search(r'\w$', term) else ''
                highlight_patterns.append(prefix + re.escape(term) + suffix)
            else:
                highlight_patterns.append(re.escape(term))

    # Combine into one pattern for efficient matching. Wrap each
    # individual highlight pattern in a non-capturing group so the
    # OR-joined pattern stays a clean alternation; otherwise capturing
    # groups inside any user pattern (e.g. (TODO|FIXME) from regex
    # collections) make re.split return interleaved capture values
    # and re.findall return tuples of captures — the kind of mismatch
    # that previously dropped words from the rendered docx. We use
    # re.finditer + span() below, which sidesteps both quirks
    # regardless of capturing structure.
    #
    # Case sensitivity is decided per-pattern with an inline (?i:...)
    # scope rather than a global re.IGNORECASE flag. Patterns whose
    # character classes explicitly use ONLY one case (e.g.
    # \b[A-Z][A-Z0-9_]{3,}\b for UPPER_CASE constants) preserve the
    # author's case intent; patterns with no case-bearing class, or
    # with both [A-Z] and [a-z], get IGNORECASE so that a literal
    # "TODO" pattern still highlights "todo" in the report (matching
    # how the search itself runs).
    def _pattern_wants_case_sensitive(pat):
        has_upper = "A-Z" in pat
        has_lower = "a-z" in pat
        # Author wrote a one-sided letter range → respect that intent.
        return has_upper != has_lower

    combined_pattern = None
    if highlight_patterns:
        try:
            parts = []
            for p in highlight_patterns:
                if _pattern_wants_case_sensitive(p):
                    parts.append(f"(?:{p})")
                else:
                    parts.append(f"(?i:{p})")
            combined_pattern = re.compile("|".join(parts))
        except re.error:
            combined_pattern = None

    def _add_highlighted_line(para, line, bold=False):
        """Add a line to a paragraph with yellow highlighting on matched terms.

        Uses re.finditer + span() rather than split/findall so capturing
        groups inside any user-supplied highlight pattern don't shift the
        match alignment. Non-matched stretches of *line* are added as
        plain runs; matched stretches get the yellow background.
        """
        if not combined_pattern:
            run = para.add_run(line)
            if bold:
                run.bold = True
            return
        last_end = 0
        added_any = False
        for m in combined_pattern.finditer(line):
            start, end = m.span()
            if start == end:
                # Zero-width match (e.g. a stray \b). Skip — adding it as
                # a highlight produces an empty run and the next iteration
                # would loop forever.
                continue
            if start > last_end:
                run = para.add_run(line[last_end:start])
                if bold:
                    run.bold = True
                added_any = True
            run = para.add_run(line[start:end])
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            if bold:
                run.bold = True
            added_any = True
            last_end = end
        if last_end < len(line):
            run = para.add_run(line[last_end:])
            if bold:
                run.bold = True
            added_any = True
        if not added_any:
            # Line had no matches and was empty after the loop (or the
            # only matches were zero-width). Emit the whole line so the
            # paragraph isn't blank.
            run = para.add_run(line)
            if bold:
                run.bold = True

    result_doc = Document()
    in_results = False
    with open(txt_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.rstrip("\n")
            para = result_doc.add_paragraph()

            # Make URL a clickable hyperlink
            if line.startswith("Program Source: "):
                prefix = "Program Source: "
                url = line[len(prefix):]
                para.add_run(prefix)
                r_id = result_doc.part.relate_to(
                    url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
                )
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("r:id"), r_id)
                run_elem = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                color = OxmlElement("w:color")
                color.set(qn("w:val"), "0000FF")
                rPr.append(color)
                u = OxmlElement("w:u")
                u.set(qn("w:val"), "single")
                rPr.append(u)
                run_elem.append(rPr)
                text_elem = OxmlElement("w:t")
                text_elem.text = url
                run_elem.append(text_elem)
                hyperlink.append(run_elem)
                para._p.append(hyperlink)
                continue

            if line.startswith("Translation ==> "):
                prefix = "Translation ==> "
                rest = line[len(prefix):]
                run = para.add_run(prefix)
                run.bold = True
                run.font.size = Pt(13)
                run = para.add_run(rest)
                run.bold = True
                run.font.size = Pt(13)
                continue

            if line.startswith("Search Term(s) ==> "):
                prefix = "Search Term(s) ==> "
                rest = line[len(prefix):]
                match = re.match(r"(.+?)( \(match: [A-Z+]+\))$", rest)
                if match:
                    terms_str, mode_str = match.group(1), match.group(2)
                    run = para.add_run(prefix)
                    run.bold = True
                    run.font.size = Pt(13)
                    run = para.add_run(terms_str)
                    run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    run.bold = True
                    run.font.size = Pt(13)
                    run = para.add_run(mode_str)
                    run.bold = True
                    run.font.size = Pt(13)
                else:
                    run = para.add_run(line)
                    run.bold = True
                    run.font.size = Pt(13)
                continue

            if line in ("Header:", "Results:"):
                run = para.add_run(line)
                run.bold = True
                in_results = (line == "Results:")
                continue

            is_doc_line = line.startswith("Document:")
            if in_results and not is_doc_line and not line.startswith("(") and not line.startswith("Files WITHOUT"):
                # Match line in results section — apply highlighting
                _add_highlighted_line(para, line)
            elif is_doc_line:
                _add_highlighted_line(para, line, bold=True)
            else:
                para.add_run(line)

    result_doc.save(docx_path)
    _restrict_file_permissions(docx_path)
    return result_doc


def insert_file_sizes(txt_path: str, docx_path: str | None, result_doc: Any) -> tuple[int, int]:
    """Insert report file sizes into both txt and docx reports.

    Returns ``(txt_size, docx_size)`` in bytes. When DOCX generation
    was skipped (``--no-docx``), pass ``docx_path=None`` and
    ``result_doc=None`` — the docx update is skipped, the sizes line in
    the TXT report mentions only the TXT, and ``docx_size`` returns 0.
    """
    txt_size = os.path.getsize(txt_path)
    txt_name = os.path.basename(txt_path)
    if docx_path is None or result_doc is None:
        docx_size = 0
        sizes_line = f"Report File Sizes ==> {txt_name} ({fmt_size(txt_size)})"
    else:
        docx_size = os.path.getsize(docx_path)
        docx_name = os.path.basename(docx_path)
        sizes_line = f"Report File Sizes ==> {txt_name} ({fmt_size(txt_size)}), {docx_name} ({fmt_size(docx_size)})"

    # Update txt report
    with open(txt_path, "r", encoding="utf-8") as f:
        content = f.read()
    content = content.replace(
        "\nCommand ==>",
        f"\n{sizes_line}\nCommand ==>",
        1,
    )
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(content)

    # Update docx report — insert sizes paragraph after timestamp
    if docx_path is not None and result_doc is not None:
        for i, para in enumerate(result_doc.paragraphs):
            if para.text.startswith("Report Generated On ==>"):
                new_para = OxmlElement("w:p")
                run_elem = OxmlElement("w:r")
                text_elem = OxmlElement("w:t")
                text_elem.text = sizes_line
                run_elem.append(text_elem)
                new_para.append(run_elem)
                para._p.addnext(new_para)
                break
        result_doc.save(docx_path)

    return (txt_size, docx_size)


def write_csv_report(
    output_path: str,
    matches: list[Match],
    inverse_files: list[str] | None = None,
) -> None:
    """Write the .csv result report."""
    if os.path.exists(output_path):
        os.remove(output_path)
    # utf-8-sig writes a UTF-8 BOM so legacy Excel on Windows reads
    # accented characters correctly instead of as cp1252 mojibake.
    # No effect on LibreOffice, Excel 365, or any text-based tooling.
    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        if inverse_files is not None:
            writer.writerow(["filename", "folder"])
            for fp in inverse_files:
                writer.writerow([os.path.basename(fp), os.path.dirname(fp)])
        else:
            writer.writerow(["filename", "folder", "line_number", "matched_text"])
            for file_dir, filename, line_num, text in matches:
                writer.writerow([filename, file_dir, line_num, _strip_highlights(text)])
    _restrict_file_permissions(output_path)


def write_json_report(
    output_path: str,
    matches: list[Match],
    search_terms: list[str],
    report_mode: str,
    files_count: int,
    search_elapsed: float,
    inverse_files: list[str] | None = None,
    directory: str | None = None,
    compute_hashes: bool = False,
) -> None:
    """Write the .json result report.

    If *compute_hashes* is True, each entry in ``matches_per_file`` /
    ``inverse_files`` gains a ``sha256`` field with the hex digest of the
    file's raw bytes (or null on read failure).
    """
    if os.path.exists(output_path):
        os.remove(output_path)

    from peekdocs.cli import VERSION as _ver_j
    if inverse_files is not None:
        inverse_entries: list[dict[str, Any]] = [
            {
                "filename": os.path.basename(fp),
                "folder": os.path.dirname(fp),
            }
            for fp in inverse_files
        ]
        if compute_hashes:
            for entry, fp in zip(inverse_entries, inverse_files):
                entry["sha256"] = _sha256_of_file(fp)
        json_data = {
            "generator": f"peekdocs v{_ver_j}",
            **({"directory": directory} if directory else {}),
            "search_terms": search_terms,
            "mode": report_mode,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "files_searched": files_count,
            "files_without_matches": len(inverse_files),
            "elapsed_seconds": round(search_elapsed, 2),
            "inverse_files": inverse_entries,
        }
    else:
        # Per-file match counts (ordered by first appearance)
        file_counts = {}
        for file_dir, filename, _ln, _tx in matches:
            key = (file_dir, filename)
            if key not in file_counts:
                file_counts[key] = 0
            file_counts[key] += 1
        matches_per_file: list[dict[str, Any]] = [
            {"filename": fn, "folder": fd, "matches": count}
            for (fd, fn), count in file_counts.items()
        ]
        if compute_hashes:
            for entry in matches_per_file:
                entry["sha256"] = _sha256_of_file(
                    os.path.join(entry["folder"], entry["filename"])
                )

        json_data = {
            "generator": f"peekdocs v{_ver_j}",
            **({"directory": directory} if directory else {}),
            "search_terms": search_terms,
            "mode": report_mode,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "files_searched": files_count,
            "matches_found": len(matches),
            "matches_per_file": matches_per_file,
            "elapsed_seconds": round(search_elapsed, 2),
            "matches": [
                {
                    "filename": filename,
                    "folder": file_dir,
                    "line_number": line_num,
                    "matched_text": _strip_highlights(text),
                }
                for file_dir, filename, line_num, text in matches
            ],
        }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(json_data, f, indent=2, ensure_ascii=False)
    _restrict_file_permissions(output_path)


def _pdf_safe(text: str) -> str:
    """Replace Unicode characters that Helvetica can't render."""
    replacements = {
        '\u2018': "'", '\u2019': "'",   # smart single quotes
        '\u201c': '"', '\u201d': '"',   # smart double quotes
        '\u2013': '-', '\u2014': '--',  # en dash, em dash
        '\u2026': '...', '\u00a0': ' ', # ellipsis, nbsp
        '\u2022': '*', '\u2023': '>',   # bullets
        '\ufb01': 'fi', '\ufb02': 'fl', # ligatures
        '\u00b0': 'deg',               # degree sign
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Remove any remaining non-latin1 characters
    return text.encode('latin-1', errors='replace').decode('latin-1')


def _pdf_insert_highlighted(pdf: Any, text: str, search_terms: list[str], highlight_re: Any = None) -> None:
    """Insert text with search terms highlighted in orange background."""
    import re as _re
    if not search_terms and not highlight_re:
        pdf.multi_cell(0, 5, text)
        return
    # Use pre-built regex if available, otherwise build from literal terms
    if highlight_re:
        pattern = highlight_re
    else:
        patterns = []
        for term in search_terms:
            patterns.append(_re.escape(term))
        try:
            pattern = _re.compile("|".join(patterns), _re.IGNORECASE)
        except _re.error:
            pdf.multi_cell(0, 5, text)
            return
    parts = pattern.split(text)
    found = pattern.findall(text)
    if not found:
        pdf.multi_cell(0, 5, text)
        return
    # Use cell for inline highlighting (multi_cell doesn't support mixed colors)
    x_start = pdf.get_x()
    page_width = pdf.w - pdf.r_margin - x_start
    for i, part in enumerate(parts):
        if part:
            pdf.set_text_color(0, 0, 0)
            pdf.write(5, part)
        if i < len(found):
            # Yellow background highlight
            pdf.set_fill_color(255, 255, 0)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(pdf.get_string_width(found[i]) + 1, 5, found[i], fill=True)
            pdf.set_fill_color(255, 255, 255)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(5)


def write_pdf_report(
    output_path: str,
    matches: list[Match],
    search_terms: list[str] | None = None,
    report_mode: str = "ANY",
    inverse_files: list[str] | None = None,
    use_regex: bool = False,
    use_wildcard: bool = False,
    use_whole_word: bool = False,
    use_fuzzy: bool = False,
    expression: str | None = None,
) -> None:
    """Write the .pdf result report with highlighted matches."""
    from fpdf import FPDF
    import re as _re_pdf

    # Build highlight regex — same logic as DOCX/HTML reporters
    _pdf_highlight_re = None
    if not use_fuzzy:
        _h_patterns = []
        if expression:
            try:
                from peekdocs.expr_parser import parse_expression, extract_positive_terms
                _h_terms = extract_positive_terms(parse_expression(expression))
            except Exception:
                _h_terms = search_terms or []
        else:
            _h_terms = search_terms or []
        for _t in _h_terms:
            if use_wildcard:
                _h_patterns.append(_wildcard_to_regex(_t))
            elif use_regex:
                _h_patterns.append(_t)
            elif use_whole_word:
                _pfx = r'\b' if _re_pdf.match(r'\w', _t) else ''
                _sfx = r'\b' if _re_pdf.search(r'\w$', _t) else ''
                _h_patterns.append(_pfx + _re_pdf.escape(_t) + _sfx)
            else:
                _h_patterns.append(_re_pdf.escape(_t))
        if _h_patterns:
            try:
                _pdf_highlight_re = _re_pdf.compile("|".join(_h_patterns), _re_pdf.IGNORECASE)
            except _re_pdf.error:
                _pdf_highlight_re = None

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    from peekdocs.cli import VERSION as _ver_pdf
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, f"peekdocs v{_ver_pdf}", ln=True)
    pdf.ln(3)

    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, _pdf_safe(f"Search terms: {' '.join(search_terms or [])}"), ln=True)
    pdf.cell(0, 6, f"Mode: {report_mode}", ln=True)
    pdf.cell(0, 6, f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
    pdf.ln(5)

    if inverse_files is not None:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, f"Files WITHOUT matches: {len(inverse_files)}", ln=True)
        pdf.ln(3)
        pdf.set_font("Helvetica", "", 10)
        for fp in inverse_files:
            filename = os.path.basename(fp)
            dirname = os.path.dirname(fp)
            pdf.cell(0, 5, _pdf_safe(f"  {filename}"), ln=True)
            pdf.set_text_color(128, 128, 128)
            pdf.cell(0, 5, _pdf_safe(f"  ({dirname})"), ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(2)
    else:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, f"Matches found: {len(matches)}", ln=True)
        pdf.ln(3)

        prev_file = None
        for file_dir, filename, line_num, text in matches:
            current_file = os.path.join(file_dir, filename)
            if current_file != prev_file:
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 10)
                fc = sum(1 for fd, fn, _ln, _tx in matches
                         if os.path.join(fd, fn) == current_file)
                pdf.set_text_color(26, 115, 232)
                pdf.cell(0, 6, _pdf_safe(f"Document: {filename} ({fc} match{'es' if fc != 1 else ''})"), ln=True)
                pdf.set_text_color(128, 128, 128)
                pdf.set_font("Helvetica", "", 9)
                pdf.cell(0, 5, _pdf_safe(f"({file_dir})"), ln=True)
                pdf.set_text_color(0, 0, 0)
                prev_file = current_file

            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(128, 128, 128)
            pdf.cell(15, 5, f"Line {line_num}:")
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 10)
            clean_text = _pdf_safe(_strip_highlights(text))
            if len(clean_text) > 200:
                clean_text = clean_text[:197] + "..."
            _pdf_insert_highlighted(pdf, clean_text, [_pdf_safe(t) for t in (search_terms or [])], highlight_re=_pdf_highlight_re)
            pdf.ln(1)

    pdf.output(output_path)
    _restrict_file_permissions(output_path)


def write_html_report(
    output_path: str,
    matches: list[Match],
    search_terms: list[str] | None = None,
    report_mode: str = "ANY",
    inverse_files: list[str] | None = None,
    use_regex: bool = False,
    use_wildcard: bool = False,
    use_whole_word: bool = False,
    use_fuzzy: bool = False,
    expression: str | None = None,
) -> None:
    """Write the .html result report with highlighted matches."""
    import html as html_mod
    import re as _re_html

    if os.path.exists(output_path):
        os.remove(output_path)

    from peekdocs.cli import VERSION as _ver_h
    terms = search_terms or []

    # Build highlight regex — same logic as DOCX reporter
    _highlight_re = None
    if not use_fuzzy:
        _h_patterns = []
        if expression:
            try:
                from peekdocs.expr_parser import parse_expression, extract_positive_terms
                _h_terms = extract_positive_terms(parse_expression(expression))
            except Exception:
                _h_terms = terms
        else:
            _h_terms = terms
        for _t in _h_terms:
            if use_wildcard:
                _h_patterns.append(_wildcard_to_regex(_t))
            elif use_regex:
                _h_patterns.append(_t)
            elif use_whole_word:
                _pfx = r'\b' if _re_html.match(r'\w', _t) else ''
                _sfx = r'\b' if _re_html.search(r'\w$', _t) else ''
                _h_patterns.append(_pfx + _re_html.escape(_t) + _sfx)
            else:
                _h_patterns.append(_re_html.escape(_t))
        if _h_patterns:
            try:
                _highlight_re = _re_html.compile("|".join(_h_patterns), _re_html.IGNORECASE)
            except _re_html.error:
                _highlight_re = None

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("<!DOCTYPE html>\n<html lang='en'>\n<head>\n")
        f.write("<meta charset='UTF-8'>\n")
        f.write(f"<title>peekdocs v{html_mod.escape(_ver_h)} Search Results</title>\n")
        f.write("<style>\n")
        f.write("body { font-family: -apple-system, 'Segoe UI', Helvetica, Arial, sans-serif; "
                "margin: 20px; background: #fff; color: #333; }\n")
        f.write("h1 { font-size: 1.5em; }\n")
        f.write(".meta { color: #666; font-size: 0.9em; margin-bottom: 1em; }\n")
        f.write(".file-header { background: #f0f4ff; padding: 8px 12px; margin-top: 1.2em; "
                "border-left: 4px solid #1a73e8; font-weight: bold; }\n")
        f.write(".file-dir { color: #888; font-size: 0.85em; padding-left: 16px; }\n")
        f.write(".match-row { padding: 4px 16px; border-bottom: 1px solid #eee; font-family: monospace; "
                "font-size: 0.9em; white-space: pre-wrap; }\n")
        f.write(".line-num { color: #999; margin-right: 8px; }\n")
        f.write("mark { background: #ffff00; padding: 1px 2px; }\n")
        f.write(".inverse-file { padding: 4px 16px; }\n")
        f.write("</style>\n</head>\n<body>\n")

        f.write(f"<h1>peekdocs v{html_mod.escape(_ver_h)}</h1>\n")
        f.write(f"<div class='meta'>Search terms: {html_mod.escape(' '.join(terms))}<br>\n")
        f.write(f"Mode: {html_mod.escape(report_mode)}<br>\n")
        f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>\n")

        if inverse_files is not None:
            f.write(f"<h2>Files WITHOUT matches: {len(inverse_files)}</h2>\n")
            for fp in inverse_files:
                fname = html_mod.escape(os.path.basename(fp))
                dname = html_mod.escape(os.path.dirname(fp))
                f.write(f"<div class='inverse-file'>{fname}<br>"
                        f"<span class='file-dir'>{dname}</span></div>\n")
        else:
            f.write(f"<p><strong>Matches found: {len(matches)}</strong></p>\n")

            prev_file = None
            for file_dir, filename, line_num, text in matches:
                current_file = os.path.join(file_dir, filename)
                if current_file != prev_file:
                    fc = sum(1 for fd, fn, _ln, _tx in matches
                             if os.path.join(fd, fn) == current_file)
                    f.write(f"<div class='file-header'>{html_mod.escape(filename)} "
                            f"({fc} match{'es' if fc != 1 else ''})</div>\n")
                    f.write(f"<div class='file-dir'>{html_mod.escape(file_dir)}</div>\n")
                    prev_file = current_file

                clean = html_mod.escape(_strip_highlights(text))
                # Highlight matches in the HTML output
                if _highlight_re:
                    # Work on the HTML-escaped text: find matches, wrap in <mark>
                    # We must match against the raw text and map positions to escaped text
                    raw_text = _strip_highlights(text)
                    parts = []
                    last_end = 0
                    for m in _highlight_re.finditer(raw_text):
                        parts.append(html_mod.escape(raw_text[last_end:m.start()]))
                        parts.append(f"<mark>{html_mod.escape(m.group())}</mark>")
                        last_end = m.end()
                    parts.append(html_mod.escape(raw_text[last_end:]))
                    clean = "".join(parts)
                f.write(f"<div class='match-row'>"
                        f"<span class='line-num'>{line_num}:</span>{clean}</div>\n")

        f.write("\n<hr>\n<p style='color:#999; font-size:0.8em;'>Generated by peekdocs</p>\n")
        f.write("</body>\n</html>\n")
    _restrict_file_permissions(output_path)


def append_results(append_name: str, output_dir: str, txt_path: str, docx_path: str | None) -> None:
    """Append current results to accumulated peekdocs files.

    When DOCX generation was skipped (``--no-docx``), pass
    ``docx_path=None`` — only the TXT accumulator is updated.
    """
    append_txt_path = os.path.join(output_dir, f"peekdocs_accumulated_{append_name}.txt")
    with open(txt_path, "r", encoding="utf-8") as src:
        results_content = src.read()
    with open(append_txt_path, "a", encoding="utf-8") as dst:
        dst.write(results_content)
    _restrict_file_permissions(append_txt_path)
    if docx_path is None:
        return
    append_docx_path = os.path.join(output_dir, f"peekdocs_accumulated_{append_name}.docx")
    if os.path.exists(append_docx_path):
        existing_doc = Document(append_docx_path)
        new_doc = Document(docx_path)
        body = existing_doc.element.body
        sect_pr = body.find(qn('w:sectPr'))
        for para in new_doc.paragraphs:
            new_elem = deepcopy(para._p)
            if sect_pr is not None:
                sect_pr.addprevious(new_elem)
            else:
                body.append(new_elem)
        existing_doc.save(append_docx_path)
    else:
        shutil.copy2(docx_path, append_docx_path)
    _restrict_file_permissions(append_docx_path)


# ── Suite Reports ──────────────────────────────────────────────


def write_suite_txt_report(
    output_path: str,
    suite_name: str,
    sections: list[dict[str, Any]],
) -> tuple[int, str]:
    """Write a combined TXT report for a search suite.

    *sections* is a list of dicts, each with:
        search_name, search_terms, matches, all_files, elapsed,
        report_mode, params (the saved-search params dict)
    """
    if os.path.exists(output_path):
        os.remove(output_path)

    total_matches = sum(s.get("total_match_count", len(s["matches"])) for s in sections)
    total_files = set()
    for s in sections:
        total_files.update(s["all_files"])
    total_elapsed = sum(s["elapsed"] for s in sections)

    with open(output_path, "w", encoding="utf-8") as f:
        from peekdocs.cli import VERSION as _ver_s
        f.write(f"peekdocs v{_ver_s}\n\n\n")
        f.write(f"Suite Report: {suite_name}\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Saved as: {os.path.abspath(output_path)}\n")
        total_matched_file_count = sum(s.get("matched_file_count", 0) for s in sections)
        f.write(f"Searches in suite: {len(sections)}\n")
        f.write(f"Total matches: {total_matches} in {total_matched_file_count} file(s)\n")
        f.write(f"Files searched: {len(total_files)}\n")
        f.write(f"Total time: {total_elapsed:.2f}s\n")
        f.write("\n")

        if len(sections) > 1:
            f.write("Section summary:\n")
            name_w = max(len(s["search_name"]) for s in sections)
            for i, s in enumerate(sections, 1):
                stot = s.get("total_match_count", len(s["matches"]))
                smfc = s.get("matched_file_count", 0)
                f.write(f"  {i}. {s['search_name']:<{name_w}}  {stot} match(es) in {smfc} file(s)\n")
            f.write("\n")

        for i, section in enumerate(sections, 1):
            name = section["search_name"]
            matches = section["matches"]
            all_files = section["all_files"]
            elapsed = section["elapsed"]
            terms = section["search_terms"]
            mode = section["report_mode"]

            f.write("=" * 72 + "\n")
            f.write(f"Search {i}/{len(sections)}: {name}\n")
            f.write(f"Terms: {' '.join(terms)} (mode: {mode})\n")
            mfc = section.get("matched_file_count", 0)
            section_total = section.get("total_match_count", len(matches))
            f.write(f"Found {section_total} match(es) in {mfc} file(s). Files searched: {len(all_files)}. Time: {elapsed:.2f}s\n")
            f.write("=" * 72 + "\n\n")

            if not matches:
                f.write("  (no matches)\n\n")
                continue

            prev_file = None
            for file_dir, filename, line_num, text in matches:
                current_file = os.path.join(file_dir, filename)
                wrapped = textwrap.fill(text, width=80) if "\n" not in text else text
                f.write(f'Document: {filename}, Line: {line_num}, Match:\n({file_dir})\n"{wrapped}"\n\n')

    _restrict_file_permissions(output_path)
    total_bytes = sum(os.path.getsize(fp) for fp in total_files if os.path.exists(fp))
    size_str = fmt_size(total_bytes)
    return (total_bytes, size_str)


def write_suite_docx_report(
    docx_path: str,
    txt_path: str,
    sections: list[dict[str, Any]],
) -> None:
    """Create a combined DOCX suite report with per-section highlighting.

    Reads the suite TXT report and applies yellow highlighting for each
    section's search terms.
    """
    if os.path.exists(docx_path):
        os.remove(docx_path)

    # Collect all highlight terms across all sections
    all_patterns = []
    for section in sections:
        terms = section.get("search_terms", [])
        params = section.get("params", {})
        use_regex = params.get("regex", False)
        use_wildcard = params.get("wildcard", False)
        use_whole_word = params.get("whole_word", False)
        expression = params.get("expression")

        if expression:
            try:
                from peekdocs.expr_parser import parse_expression, extract_positive_terms
                terms = extract_positive_terms(parse_expression(expression))
            except Exception:
                pass

        for term in terms:
            if use_wildcard:
                all_patterns.append(_wildcard_to_regex(term))
            elif use_regex:
                all_patterns.append(term)
            elif use_whole_word:
                prefix = r'\b' if re.match(r'\w', term) else ''
                suffix = r'\b' if re.search(r'\w$', term) else ''
                all_patterns.append(prefix + re.escape(term) + suffix)
            else:
                all_patterns.append(re.escape(term))

    combined_pattern = None
    if all_patterns:
        try:
            combined_pattern = re.compile("|".join(all_patterns), re.IGNORECASE)
        except re.error:
            combined_pattern = None

    # Read the TXT report and build the DOCX
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Consolas"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)

    with open(txt_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.rstrip("\n")
            para = doc.add_paragraph()

            # Section headers get bold formatting
            if line.startswith("Suite Report:") or line.startswith("Search ") or line.startswith("Section summary:"):
                run = para.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
                continue
            if line.startswith("=" * 10):
                run = para.add_run(line)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                continue

            # Apply yellow highlighting to matched terms
            if combined_pattern and combined_pattern.search(line):
                pos = 0
                for m in combined_pattern.finditer(line):
                    if m.start() > pos:
                        para.add_run(line[pos:m.start()])
                    run = para.add_run(m.group())
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    pos = m.end()
                if pos < len(line):
                    para.add_run(line[pos:])
            else:
                para.add_run(line)

    doc.save(docx_path)
    _restrict_file_permissions(docx_path)


def write_suite_html_report(
    output_path: str,
    suite_name: str,
    sections: list[dict[str, Any]],
) -> None:
    """Write a combined HTML report for a search suite with highlighted matches."""
    import html as html_mod
    import re as _re_sh

    if os.path.exists(output_path):
        os.remove(output_path)

    from peekdocs.cli import VERSION as _ver_sh

    # Collect highlight patterns from all sections
    all_patterns = []
    for section in sections:
        terms = section.get("search_terms", [])
        params = section.get("params", {})
        use_regex = params.get("regex", False)
        use_wildcard = params.get("wildcard", False)
        use_whole_word = params.get("whole_word", False)
        for term in terms:
            if use_wildcard:
                from peekdocs.scanner import _wildcard_to_regex
                all_patterns.append(_wildcard_to_regex(term))
            elif use_regex:
                all_patterns.append(term)
            elif use_whole_word:
                prefix = r'\b' if _re_sh.match(r'\w', term) else ''
                suffix = r'\b' if _re_sh.search(r'\w$', term) else ''
                all_patterns.append(prefix + _re_sh.escape(term) + suffix)
            else:
                all_patterns.append(_re_sh.escape(term))

    combined_re = None
    if all_patterns:
        try:
            combined_re = _re_sh.compile("|".join(all_patterns), _re_sh.IGNORECASE)
        except _re_sh.error:
            combined_re = None

    total_matches = sum(s.get("total_match_count", len(s["matches"])) for s in sections)
    total_matched_files = sum(s.get("matched_file_count", 0) for s in sections)
    total_elapsed = sum(s["elapsed"] for s in sections)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("<!DOCTYPE html>\n<html lang='en'>\n<head>\n")
        f.write("<meta charset='UTF-8'>\n")
        f.write(f"<title>peekdocs v{html_mod.escape(_ver_sh)} — Suite: {html_mod.escape(suite_name)}</title>\n")
        f.write("<style>\n")
        f.write("body { font-family: -apple-system, 'Segoe UI', Helvetica, Arial, sans-serif; "
                "margin: 20px; background: #fff; color: #333; }\n")
        f.write("h1 { font-size: 1.5em; }\n")
        f.write("h2 { font-size: 1.2em; margin-top: 1.5em; border-bottom: 2px solid #ccc; padding-bottom: 4px; }\n")
        f.write(".meta { color: #666; font-size: 0.9em; margin-bottom: 1em; }\n")
        f.write(".match { margin: 6px 0; padding: 6px 12px; background: #f8f8f8; "
                "border-left: 3px solid #2196F3; font-family: monospace; font-size: 0.9em; "
                "white-space: pre-wrap; }\n")
        f.write("mark { background: #ffff00; padding: 1px 2px; }\n")
        f.write(".section-meta { color: #888; font-size: 0.85em; }\n")
        f.write(".toc { background: #f0f7ff; border-left: 4px solid #2196F3; "
                "padding: 10px 16px; margin: 1em 0; }\n")
        f.write(".toc ol { margin: 4px 0 0 0; padding-left: 24px; }\n")
        f.write(".toc a { color: #1565C0; text-decoration: none; }\n")
        f.write(".toc a:hover { text-decoration: underline; }\n")
        f.write("</style>\n</head>\n<body>\n")

        f.write(f"<h1>peekdocs v{html_mod.escape(_ver_sh)} — Suite: {html_mod.escape(suite_name)}</h1>\n")
        f.write(f"<div class='meta'>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br>\n")
        f.write(f"Searches: {len(sections)} | Total matches: {total_matches} in {total_matched_files} file(s)<br>\n")
        f.write(f"Total time: {total_elapsed:.2f}s</div>\n")

        if len(sections) > 1:
            f.write("<div class='toc'><b>Section summary:</b>\n<ol>\n")
            for i, s in enumerate(sections, 1):
                stot = s.get("total_match_count", len(s["matches"]))
                smfc = s.get("matched_file_count", 0)
                f.write(f"<li><a href='#sec{i}'>{html_mod.escape(s['search_name'])}</a> "
                        f"— {stot} match(es) in {smfc} file(s)</li>\n")
            f.write("</ol></div>\n")

        for i, section in enumerate(sections, 1):
            name = section["search_name"]
            matches = section["matches"]
            mfc = section.get("matched_file_count", 0)
            files_searched = len(section["all_files"])
            elapsed = section["elapsed"]
            s_total = section.get("total_match_count", len(matches))

            f.write(f"<h2 id='sec{i}'>Search {i}/{len(sections)}: {html_mod.escape(name)}</h2>\n")
            f.write(f"<div class='section-meta'>Terms: {html_mod.escape(' '.join(section['search_terms']))} "
                    f"| Mode: {section['report_mode']} | "
                    f"Found {s_total} match(es) in {mfc} file(s) | "
                    f"Files searched: {files_searched} | Time: {elapsed:.2f}s</div>\n")

            if not matches:
                f.write("<p><em>(no matches)</em></p>\n")
                continue

            for fd, fn, ln, text in matches:
                safe_text = html_mod.escape(text)
                if combined_re:
                    safe_text = combined_re.sub(
                        lambda m: f"<mark>{html_mod.escape(m.group())}</mark>",
                        safe_text
                    )
                f.write(f"<div class='match'><b>{html_mod.escape(fn)}</b> line {ln}:<br>{safe_text}</div>\n")

        f.write("</body>\n</html>\n")
    _restrict_file_permissions(output_path)
